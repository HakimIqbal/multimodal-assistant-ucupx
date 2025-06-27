# pip install python-docx reportlab
from fastapi import APIRouter, Depends, HTTPException, Query
from typing import List, Optional
from datetime import datetime
from pydantic import BaseModel
from api.auth.auth_middleware import get_current_user
import json
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
import uuid
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.responses import StreamingResponse

router = APIRouter()

class ExportRequest(BaseModel):
    chat_type: str  # "general", "coder", "rag"
    session_id: str
    format: str = "pdf"  # "pdf", "docx", "txt"
    include_metadata: bool = True
    include_timestamps: bool = True

def generate_pdf_chat_export(chat_data: List[dict], metadata: dict) -> bytes:
    """Generate PDF export of chat conversation"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1
    )
    story.append(Paragraph("Chat Export", title_style))
    story.append(Spacer(1, 12))
    
    # Metadata
    if metadata.get("include_metadata", True):
        meta_style = ParagraphStyle(
            'MetaData',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey
        )
        story.append(Paragraph(f"<b>Chat Type:</b> {metadata.get('chat_type', 'Unknown')}", meta_style))
        story.append(Paragraph(f"<b>Session ID:</b> {metadata.get('session_id', 'Unknown')}", meta_style))
        story.append(Paragraph(f"<b>Export Date:</b> {metadata.get('export_date', 'Unknown')}", meta_style))
        story.append(Paragraph(f"<b>Total Messages:</b> {len(chat_data)}", meta_style))
        story.append(Spacer(1, 12))
    
    # Chat messages
    for i, message in enumerate(chat_data):
        # Message header
        header_style = ParagraphStyle(
            'MessageHeader',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=6
        )
        
        role = message.get("role", "Unknown")
        timestamp = message.get("timestamp", "")
        
        if metadata.get("include_timestamps", True) and timestamp:
            header_text = f"{role} - {timestamp}"
        else:
            header_text = role
        
        story.append(Paragraph(header_text, header_style))
        
        # Message content
        content_style = ParagraphStyle(
            'MessageContent',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=12,
            leftIndent=20
        )
        
        content = message.get("content", "")
        # Handle code blocks
        if "```" in content:
            # Simple code block handling
            content = content.replace("```", "<code>").replace("<code>", "</code>")
        
        story.append(Paragraph(content, content_style))
        
        if i < len(chat_data) - 1:
            story.append(Spacer(1, 6))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_docx_chat_export(chat_data: List[dict], metadata: dict) -> bytes:
    """Generate DOCX export of chat conversation"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Chat Export', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    if metadata.get("include_metadata", True):
        doc.add_paragraph(f"Chat Type: {metadata.get('chat_type', 'Unknown')}")
        doc.add_paragraph(f"Session ID: {metadata.get('session_id', 'Unknown')}")
        doc.add_paragraph(f"Export Date: {metadata.get('export_date', 'Unknown')}")
        doc.add_paragraph(f"Total Messages: {len(chat_data)}")
        doc.add_paragraph("")  # Empty line
    
    # Chat messages
    for message in chat_data:
        role = message.get("role", "Unknown")
        timestamp = message.get("timestamp", "")
        content = message.get("content", "")
        
        # Message header
        if metadata.get("include_timestamps", True) and timestamp:
            header_text = f"{role} - {timestamp}"
        else:
            header_text = role
        
        paragraph = doc.add_paragraph(header_text, style='List Bullet')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Message content
        paragraph = doc.add_paragraph(content)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_txt_chat_export(chat_data: List[dict], metadata: dict) -> str:
    """Generate TXT export of chat conversation"""
    lines = []
    
    # Title
    lines.append("=" * 50)
    lines.append("CHAT EXPORT")
    lines.append("=" * 50)
    lines.append("")
    
    # Metadata
    if metadata.get("include_metadata", True):
        lines.append(f"Chat Type: {metadata.get('chat_type', 'Unknown')}")
        lines.append(f"Session ID: {metadata.get('session_id', 'Unknown')}")
        lines.append(f"Export Date: {metadata.get('export_date', 'Unknown')}")
        lines.append(f"Total Messages: {len(chat_data)}")
        lines.append("")
    
    # Chat messages
    for message in chat_data:
        role = message.get("role", "Unknown")
        timestamp = message.get("timestamp", "")
        content = message.get("content", "")
        
        # Message header
        if metadata.get("include_timestamps", True) and timestamp:
            header_text = f"{role} - {timestamp}"
        else:
            header_text = role
        
        lines.append(f"[{header_text}]")
        lines.append("-" * 30)
        lines.append(content)
        lines.append("")
    
    return "\n".join(lines)

@router.post("/export/chat")
async def export_chat(request: ExportRequest, user=Depends(get_current_user)):
    """
    Export chat conversation to various formats
    """
    try:
        from src.db import supabase
        
        # Validate chat type
        valid_types = ["general", "coder", "rag"]
        if request.chat_type not in valid_types:
            raise HTTPException(status_code=400, detail=f"Invalid chat type. Valid types: {valid_types}")
        
        # Validate format
        valid_formats = ["pdf", "docx", "txt"]
        if request.format not in valid_formats:
            raise HTTPException(status_code=400, detail=f"Invalid format. Valid formats: {valid_formats}")
        
        # Get chat data from database
        table_name = f"{request.chat_type}_logs"
        res = supabase.table(table_name).select("*").eq("session_id", request.session_id).order("timestamp").execute()
        
        if not res.data:
            raise HTTPException(status_code=404, detail="No chat data found for this session")
        
        # Transform data for export
        chat_data = []
        for item in res.data:
            chat_data.append({
                "role": "User" if "input" in item else "Assistant",
                "content": item.get("input", "") or item.get("output", ""),
                "timestamp": item.get("timestamp", "")
            })
        
        # Prepare metadata
        metadata = {
            "chat_type": request.chat_type,
            "session_id": request.session_id,
            "export_date": datetime.utcnow().isoformat(),
            "include_metadata": request.include_metadata,
            "include_timestamps": request.include_timestamps
        }
        
        # Generate export based on format
        if request.format == "pdf":
            content = generate_pdf_chat_export(chat_data, metadata)
            content_type = "application/pdf"
            filename = f"chat_export_{request.chat_type}_{request.session_id[:8]}.pdf"
        elif request.format == "docx":
            content = generate_docx_chat_export(chat_data, metadata)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"chat_export_{request.chat_type}_{request.session_id[:8]}.docx"
        else:  # txt
            content = generate_txt_chat_export(chat_data, metadata)
            content_type = "text/plain"
            filename = f"chat_export_{request.chat_type}_{request.session_id[:8]}.txt"
        
        # Log export activity
        export_log = {
            "id": str(uuid.uuid4()),
            "user_id": user["id"],
            "chat_type": request.chat_type,
            "session_id": request.session_id,
            "format": request.format,
            "export_date": datetime.utcnow().isoformat(),
            "message_count": len(chat_data)
        }
        
        # Save export log (implement this table)
        # supabase.table("export_logs").insert(export_log).execute()
        
        return {
            "success": True,
            "filename": filename,
            "format": request.format,
            "message_count": len(chat_data),
            "export_date": metadata["export_date"],
            "download_url": f"/api/export/download/{export_log['id']}"  # Implement download endpoint
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to export chat: {str(e)}")

@router.post("/export/document")
async def export_document_analysis(
    document_id: str,
    format: str = "pdf",
    include_summary: bool = True,
    include_qa_history: bool = True,
    user=Depends(get_current_user)
):
    """
    Export document analysis and Q&A history
    """
    try:
        from src.db import supabase
        
        # Get document
        doc_res = supabase.table("documents").select("*").eq("id", document_id).eq("user_id", user["id"]).execute()
        
        if not doc_res.data:
            raise HTTPException(status_code=404, detail="Document not found")
        
        document = doc_res.data[0]
        
        # Get Q&A history
        qa_history = []
        if include_qa_history:
            qa_res = supabase.table("rag_logs").select("*").eq("metadata->source_info->filename", document["filename"]).order("timestamp").execute()
            qa_history = qa_res.data or []
        
        # Generate export
        if format == "pdf":
            content = generate_document_pdf_export(document, qa_history, include_summary)
            content_type = "application/pdf"
            filename = f"document_analysis_{document['filename']}.pdf"
        elif format == "docx":
            content = generate_document_docx_export(document, qa_history, include_summary)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"document_analysis_{document['filename']}.docx"
        else:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX formats supported for document export")
        
        return {
            "success": True,
            "filename": filename,
            "format": format,
            "document_name": document["filename"],
            "qa_count": len(qa_history),
            "export_date": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to export document: {str(e)}")

def generate_document_pdf_export(document: dict, qa_history: List[dict], include_summary: bool) -> bytes:
    """Generate PDF export of document analysis"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1
    )
    story.append(Paragraph("Document Analysis Export", title_style))
    story.append(Spacer(1, 12))
    
    # Document info
    info_style = ParagraphStyle(
        'Info',
        parent=styles['Normal'],
        fontSize=10
    )
    story.append(Paragraph(f"<b>Document:</b> {document.get('filename', 'Unknown')}", info_style))
    story.append(Paragraph(f"<b>Category:</b> {document.get('category', 'Unknown')}", info_style))
    story.append(Paragraph(f"<b>Upload Date:</b> {document.get('upload_timestamp', 'Unknown')}", info_style))
    story.append(Paragraph(f"<b>Word Count:</b> {document.get('word_count', 0)}", info_style))
    story.append(Spacer(1, 12))
    
    # Summary
    if include_summary and document.get("summary"):
        story.append(Paragraph("Document Summary", styles['Heading2']))
        story.append(Paragraph(document["summary"], styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Key topics
    if document.get("key_topics"):
        story.append(Paragraph("Key Topics", styles['Heading2']))
        topics_text = ", ".join(document["key_topics"])
        story.append(Paragraph(topics_text, styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Q&A History
    if qa_history:
        story.append(Paragraph("Q&A History", styles['Heading2']))
        
        for qa in qa_history:
            question = qa.get("input", "")
            answer = qa.get("output", "")
            timestamp = qa.get("timestamp", "")
            
            if question and answer:
                story.append(Paragraph(f"<b>Q:</b> {question}", styles['Normal']))
                story.append(Paragraph(f"<b>A:</b> {answer}", styles['Normal']))
                if timestamp:
                    story.append(Paragraph(f"<i>Time: {timestamp}</i>", styles['Normal']))
                story.append(Spacer(1, 6))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_document_docx_export(document: dict, qa_history: List[dict], include_summary: bool) -> bytes:
    """Generate DOCX export of document analysis"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Document Analysis Export', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document info
    doc.add_paragraph(f"Document: {document.get('filename', 'Unknown')}")
    doc.add_paragraph(f"Category: {document.get('category', 'Unknown')}")
    doc.add_paragraph(f"Upload Date: {document.get('upload_timestamp', 'Unknown')}")
    doc.add_paragraph(f"Word Count: {document.get('word_count', 0)}")
    doc.add_paragraph("")
    
    # Summary
    if include_summary and document.get("summary"):
        doc.add_heading('Document Summary', level=1)
        doc.add_paragraph(document["summary"])
        doc.add_paragraph("")
    
    # Key topics
    if document.get("key_topics"):
        doc.add_heading('Key Topics', level=1)
        topics_text = ", ".join(document["key_topics"])
        doc.add_paragraph(topics_text)
        doc.add_paragraph("")
    
    # Q&A History
    if qa_history:
        doc.add_heading('Q&A History', level=1)
        
        for qa in qa_history:
            question = qa.get("input", "")
            answer = qa.get("output", "")
            timestamp = qa.get("timestamp", "")
            
            if question and answer:
                doc.add_paragraph(f"Q: {question}")
                doc.add_paragraph(f"A: {answer}")
                if timestamp:
                    doc.add_paragraph(f"Time: {timestamp}")
                doc.add_paragraph("")
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

@router.get("/export/stats")
async def get_export_stats(user=Depends(get_current_user)):
    """
    Get export statistics for user
    """
    try:
        from src.db import supabase
        res = supabase.table("export_logs").select("*").eq("user_id", user["id"]).execute()
        data = res.data if res and res.data else []
        total_exports = len(data)
        exports_by_format = {"pdf": 0, "docx": 0, "txt": 0}
        exports_by_type = {"chat": 0, "document": 0}
        for d in data:
            fmt = d.get("format", "pdf")
            typ = d.get("type", "chat")
            exports_by_format[fmt] = exports_by_format.get(fmt, 0) + 1
            exports_by_type[typ] = exports_by_type.get(typ, 0) + 1
        recent_exports = data[-10:]
        return {
            "total_exports": total_exports,
            "exports_by_format": exports_by_format,
            "exports_by_type": exports_by_type,
            "recent_exports": recent_exports,
            "user_id": user["id"]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get export stats: {str(e)}")

@router.get("/export/download/{id}")
async def download_exported_file(id: str, user=Depends(get_current_user)):
    """
    Download exported file (PDF/DOCX/TXT) by export log ID
    """
    try:
        from src.db import supabase
        # Cari log export
        log_res = supabase.table("export_logs").select("*").eq("id", id).eq("user_id", user["id"]).execute()
        if not log_res.data:
            raise HTTPException(status_code=404, detail="Export log not found")
        export_log = log_res.data[0]
        filename = export_log.get("filename")
        file_format = export_log.get("format")
        # Ambil file dari storage (misal supabase storage atau local)
        # Contoh: file disimpan di supabase storage bucket 'exports'
        file_path = f"exports/{user['id']}/{filename}"
        file_data = supabase.storage.from_("exports").download(file_path)
        if not file_data:
            raise HTTPException(status_code=404, detail="Exported file not found in storage")
        # Tentukan content type
        if file_format == "pdf":
            content_type = "application/pdf"
        elif file_format == "docx":
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            content_type = "text/plain"
        return StreamingResponse(io.BytesIO(file_data), media_type=content_type, headers={
            "Content-Disposition": f"attachment; filename={filename}"
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to download exported file: {str(e)}") 